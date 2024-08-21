import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import random
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


def read_question_bank_csv(filename):
    df = pd.read_csv(filename)
    topics = {}

    for _, row in df.iterrows():
        topic = row["Topic"]
        question_text = row["QuestionText"]
        marks = row["Marks"]
        image_path = row["ImagePath"] if not pd.isna(row["ImagePath"]) else None

        if topic not in topics:
            topics[topic] = {}
        if marks not in topics[topic]:
            topics[topic][marks] = []
        topics[topic][marks].append((question_text, image_path, [], marks))

    return topics


def set_document_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Verdana"
    font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def select_questions(questions_list, num_questions):
    if num_questions <= len(questions_list):
        return random.sample(questions_list, num_questions)
    else:
        return questions_list + random.choices(
            questions_list, k=num_questions - len(questions_list)
        )


def add_image_to_document(doc, image_path):
    doc.add_picture(image_path, width=Inches(4.0))


def create_question_paper(
    topics, num_questions_per_type, specific_topic_questions, output_filename
):
    doc = Document()
    set_document_style(doc)

    doc.add_heading("Question Paper", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    selected_questions = []

    for topic, num_questions_by_marks in specific_topic_questions.items():
        for marks, num_questions in num_questions_by_marks.items():
            if marks in num_questions_per_type:
                topic_questions = topics.get(topic, {}).get(marks, [])
                selected_questions.extend(
                    select_questions(topic_questions, num_questions)
                )

    remaining_questions_per_type = {
        mark: count - sum(q[3] == mark for q in selected_questions)
        for mark, count in num_questions_per_type.items()
    }

    if any(count > 0 for count in remaining_questions_per_type.values()):
        for marks, num_questions in remaining_questions_per_type.items():
            if num_questions > 0:
                all_topic_questions = []
                for topic, topic_questions in topics.items():
                    if topic != "Current Electricity":
                        all_topic_questions.extend(topic_questions.get(marks, []))
                selected_questions.extend(
                    select_questions(all_topic_questions, num_questions)
                )

    questions_by_marks = {}
    for question in selected_questions:
        marks = question[3]
        if marks not in questions_by_marks:
            questions_by_marks[marks] = []
        questions_by_marks[marks].append(question)

    for marks, questions in questions_by_marks.items():
        doc.add_heading(f"{marks}-Mark Questions", level=1)
        for question_text, image_path, _, _ in questions:
            doc.add_paragraph(f"Question: {question_text}")
            if image_path:
                add_image_to_document(doc, image_path)

    doc.save(output_filename)


class QuestionPaperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Question Paper Generator")
        self.root.geometry("600x400")

        self.filename = ""
        self.topics = {}
        self.num_questions_per_type = {}
        self.specific_topic_questions = {}

        self.set_theme()
        self.create_widgets()

    def set_theme(self):
        # Define the dark color scheme
        self.primary_color = "#2E2E2E"
        self.secondary_color = "#FF8C00"
        self.text_color = "white"

        style = ttk.Style()
        style.theme_use("clam")

        # Configure the theme
        style.configure("TFrame", background=self.primary_color)
        style.configure(
            "TLabel",
            background=self.primary_color,
            foreground=self.text_color,
            font=("Helvetica", 10),
        )
        style.configure(
            "TButton",
            background=self.primary_color,
            foreground=self.text_color,
            font=("Helvetica", 10, "bold"),
            highlightbackground=self.secondary_color,
        )
        style.map(
            "TButton",
            background=[("active", self.secondary_color)],
            foreground=[("active", self.text_color)],
        )
        style.configure(
            "TEntry",
            fieldbackground=self.primary_color,
            foreground=self.text_color,
            insertcolor=self.text_color,
        )

    def create_widgets(self):
        self.select_file_button = ttk.Button(
            self.root, text="Select Question Bank CSV", command=self.load_csv
        )
        self.select_file_button.pack(pady=10)

        self.questions_frame = ttk.Frame(self.root)
        self.questions_frame.pack(fill="x", pady=10)

        self.question_labels = {}
        self.question_entries = {}
        for mark_type in [1, 2, 3, 5]:
            lbl = ttk.Label(
                self.questions_frame, text=f"Number of {mark_type}-mark questions:"
            )
            lbl.grid(row=mark_type - 1, column=0, padx=5, pady=5, sticky="e")

            entry = ttk.Entry(self.questions_frame, width=10)
            entry.grid(row=mark_type - 1, column=1, padx=5, pady=5, sticky="w")

            self.question_labels[mark_type] = lbl
            self.question_entries[mark_type] = entry

        self.generate_button = ttk.Button(
            self.root,
            text="Generate Question Paper",
            command=self.generate_question_paper,
        )
        self.generate_button.pack(pady=20)

    def load_csv(self):
        self.filename = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
        if self.filename:
            self.topics = read_question_bank_csv(self.filename)
            messagebox.showinfo("Success", "Question bank loaded successfully!")

    def generate_question_paper(self):
        output_filename = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word Document", "*.docx")]
        )

        if not output_filename:
            return

        try:
            for mark_type in [1, 2, 3, 5]:
                num_questions = int(self.question_entries[mark_type].get())
                self.num_questions_per_type[mark_type] = num_questions

            self.specific_topic_questions = {topic: {} for topic in self.topics}

            for topic in self.topics:
                for mark_type in self.num_questions_per_type:
                    max_questions = self.num_questions_per_type[mark_type]
                    num_questions = min(
                        max_questions, len(self.topics[topic].get(mark_type, []))
                    )
                    self.specific_topic_questions[topic][mark_type] = num_questions

            create_question_paper(
                self.topics,
                self.num_questions_per_type,
                self.specific_topic_questions,
                output_filename,
            )
            messagebox.showinfo("Success", "Question paper generated successfully!")

        except ValueError:
            messagebox.showerror(
                "Error", "Please enter valid numbers for all question types."
            )


if __name__ == "__main__":
    root = tk.Tk()
    app = QuestionPaperApp(root)
    root.mainloop()
